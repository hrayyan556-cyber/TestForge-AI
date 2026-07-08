from __future__ import annotations

import io
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pypdf import PdfReader
from sqlmodel import Session

from app.database import get_session
from app.models import Project, UploadedDocument, User
from app.schemas import UploadExtractResponse
from app.utils.security import get_current_user

router = APIRouter(prefix="/api/uploads", tags=["uploads"])


async def _extract_text_from_upload(file: UploadFile) -> str:
    raw = await file.read()
    filename = (file.filename or "uploaded-file").lower()
    content_type = (file.content_type or "").lower()

    if filename.endswith(".pdf") or "pdf" in content_type:
        try:
            reader = PdfReader(io.BytesIO(raw))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(page.strip() for page in pages if page.strip())
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Could not read PDF: {str(exc)[:120]}") from exc

    if filename.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Could not read DOCX: {str(exc)[:120]}") from exc

    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return raw.decode("latin-1")
        except Exception as exc:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, TXT, or MD.") from exc


def _suggest_title(text: str, filename: str) -> str:
    for line in text.splitlines():
        clean = line.strip(" #\t-—")
        if 8 <= len(clean) <= 120:
            return clean[:120]
    return (filename.rsplit(".", 1)[0] or "Uploaded Requirement")[:120]


@router.post("/extract-requirement", response_model=UploadExtractResponse)
async def extract_requirement(
    project_id: Optional[int] = Form(default=None),
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    if project_id is not None:
        project = session.get(Project, project_id)
        if not project or project.user_id != current_user.id:
            raise HTTPException(status_code=404, detail="Project not found")

    text = (await _extract_text_from_upload(file)).strip()
    if len(text) < 10:
        raise HTTPException(status_code=400, detail="Could not extract enough readable text from this file")

    max_chars = 12000
    clipped = text[:max_chars]
    doc = UploadedDocument(
        user_id=current_user.id,
        project_id=project_id,
        filename=file.filename or "uploaded-file",
        content_type=file.content_type or "",
        extracted_text=clipped,
    )
    session.add(doc)
    session.commit()

    return UploadExtractResponse(
        filename=file.filename or "uploaded-file",
        character_count=len(clipped),
        suggested_title=_suggest_title(clipped, file.filename or "Uploaded Requirement"),
        extracted_text=clipped,
    )
